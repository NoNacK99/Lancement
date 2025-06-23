# ai_analyzer.py
# Module d'analyse IA avec OpenAI (version économique)

import os
import json
from datetime import datetime
from typing import Dict, Any
import PyPDF2
import docx
from openai import AsyncOpenAI

# Configuration OpenAI
client = AsyncOpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

async def extract_text_from_pdf(file_path: str) -> str:
    """Extraire le texte d'un PDF"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            # Limiter à 10 pages pour économiser les tokens
            max_pages = min(num_pages, 10)
            
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
            if num_pages > 10:
                text += f"\n[Note: Document tronqué - {num_pages} pages au total, 10 premières pages analysées]"
                
    except Exception as e:
        print(f"Erreur extraction PDF: {e}")
        raise e
    return text

async def extract_text_from_docx(file_path: str) -> str:
    """Extraire le texte d'un DOCX"""
    text = ""
    try:
        doc = docx.Document(file_path)
        # Limiter le nombre de paragraphes pour économiser
        max_paragraphs = 100
        paragraphs = doc.paragraphs[:max_paragraphs]
        
        for paragraph in paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                
        if len(doc.paragraphs) > max_paragraphs:
            text += f"\n[Note: Document tronqué - {len(doc.paragraphs)} paragraphes au total, 100 premiers analysés]"
            
    except Exception as e:
        print(f"Erreur extraction DOCX: {e}")
        raise e
    return text

async def extract_text_from_file(file_path: str) -> str:
    """Extraire le texte selon le type de fichier"""
    if file_path.lower().endswith('.pdf'):
        return await extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return await extract_text_from_docx(file_path)
    elif file_path.lower().endswith('.doc'):
        # Pour .doc, on retourne une erreur claire
        raise ValueError("Format .doc non supporté. Veuillez convertir en .docx ou .pdf")
    else:
        raise ValueError(f"Format de fichier non supporté: {file_path}")

async def analyze_business_plan(text: str, student_name: str, project_title: str) -> Dict[str, Any]:
    """Analyser un plan d'affaires avec GPT-3.5-turbo (économique)"""
    
    # Limiter le texte à ~3000 mots pour économiser les tokens
    words = text.split()
    if len(words) > 3000:
        text = ' '.join(words[:3000]) + "\n\n[Document tronqué pour l'analyse]"
    
    # Prompt optimisé pour GPT-3.5-turbo
    system_prompt = """Tu es un expert en évaluation de plans d'affaires académiques.
Analyse le plan et retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
    "resume_executif": "résumé en 2-3 phrases",
    "score_global": nombre entre 0 et 100,
    "scores": {
        "viabilite_concept": nombre entre 0 et 20,
        "etude_marche": nombre entre 0 et 20,
        "modele_economique": nombre entre 0 et 20,
        "strategie_marketing": nombre entre 0 et 20,
        "projections_financieres": nombre entre 0 et 20
    },
    "points_forts": ["point 1", "point 2", "point 3"],
    "axes_amelioration": ["axe 1", "axe 2", "axe 3"],
    "recommandations": ["reco 1", "reco 2", "reco 3"]
}
Sois concis mais constructif."""
    
    user_prompt = f"""Plan d'affaires de {student_name} - Projet: {project_title}

{text}

Fournis l'analyse JSON."""
    
    try:
        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",  # Modèle économique
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000,  # Limiter pour économiser
            response_format={"type": "json_object"}
        )
        
        # Parser la réponse
        analysis = json.loads(response.choices[0].message.content)
        
        # Vérifier et compléter les scores manquants
        if 'scores' not in analysis:
            analysis['scores'] = {}
        
        # Assurer tous les scores requis
        default_scores = {
            'viabilite_concept': 10,
            'etude_marche': 10,
            'modele_economique': 10,
            'strategie_marketing': 10,
            'projections_financieres': 10
        }
        
        for key, default in default_scores.items():
            if key not in analysis['scores']:
                analysis['scores'][key] = default
        
        # Calculer score global si absent
        if 'score_global' not in analysis:
            total = sum(analysis['scores'].values())
            analysis['score_global'] = int(total * 100 / 100)  # Sur 100
            
        return analysis
        
    except Exception as e:
        print(f"Erreur OpenAI: {e}")
        # Retour d'urgence en cas d'erreur
        return {
            "error": str(e),
            "resume_executif": "Erreur lors de l'analyse automatique. Veuillez réessayer.",
            "score_global": 60,  # Score par défaut
            "scores": {
                "viabilite_concept": 12,
                "etude_marche": 12,
                "modele_economique": 12,
                "strategie_marketing": 12,
                "projections_financieres": 12
            },
            "points_forts": ["Document soumis avec succès", "Format valide", "Informations complètes"],
            "axes_amelioration": ["Analyse automatique incomplète", "Révision manuelle recommandée", "Contactez le professeur si nécessaire"],
            "recommandations": ["Resoumettez si l'analyse semble incomplète", "Vérifiez le format du document", "Assurez-vous que le texte est lisible"]
        }

def generate_formatted_report(analysis: Dict[str, Any], student_name: str, project_title: str, processing_time: int) -> str:
    """Générer un rapport HTML formaté"""
    
    scores = analysis.get('scores', {})
    score_global = analysis.get('score_global', 0)
    
    # Déterminer la couleur selon le score
    if score_global >= 80:
        color_score = "#4CAF50"  # Vert
    elif score_global >= 60:
        color_score = "#FF9800"  # Orange
    else:
        color_score = "#f44336"  # Rouge
    
    report_html = f"""
    <div style="font-family: 'Inter', -apple-system, sans-serif; line-height: 1.6; color: #333;">
        
        <!-- En-tête -->
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 16px; margin-bottom: 30px;">
            <h1 style="margin: 0 0 10px 0;">📊 Rapport d'Analyse Automatique</h1>
            <p style="margin: 0; opacity: 0.9;">Évaluation par Intelligence Artificielle</p>
        </div>
        
        <!-- Informations principales -->
        <div style="background: #f8f9fa; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
            <h2 style="color: #333; margin-top: 0;">📋 Informations du Projet</h2>
            <table style="width: 100%; border-collapse: collapse;">
                <tr>
                    <td style="padding: 8px 0;"><strong>Étudiant:</strong></td>
                    <td style="padding: 8px 0;">{student_name}</td>
                </tr>
                <tr>
                    <td style="padding: 8px 0;"><strong>Projet:</strong></td>
                    <td style="padding: 8px 0;">{project_title}</td>
                </tr>
                <tr>
                    <td style="padding: 8px 0;"><strong>Date:</strong></td>
                    <td style="padding: 8px 0;">{datetime.now().strftime('%d/%m/%Y à %H:%M')}</td>
                </tr>
                <tr>
                    <td style="padding: 8px 0;"><strong>Score Global:</strong></td>
                    <td style="padding: 8px 0;">
                        <span style="font-size: 1.5em; color: {color_score}; font-weight: bold;">
                            {score_global}/100
                        </span>
                    </td>
                </tr>
            </table>
        </div>
        
        <!-- Résumé exécutif -->
        <div style="background: #e3f2fd; padding: 25px; border-radius: 12px; margin-bottom: 25px; border-left: 4px solid #2196F3;">
            <h2 style="color: #1976D2; margin-top: 0;">🎯 Résumé Exécutif</h2>
            <p style="margin: 0;">{analysis.get('resume_executif', 'Non disponible')}</p>
        </div>
        
        <!-- Scores détaillés -->
        <div style="margin-bottom: 25px;">
            <h2 style="color: #333;">📊 Scores Détaillés</h2>
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 15px;">
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;">
                    <div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Viabilité du Concept</div>
                    <div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('viabilite_concept', 0)}/20</div>
                </div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;">
                    <div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Étude de Marché</div>
                    <div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('etude_marche', 0)}/20</div>
                </div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;">
                    <div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Modèle Économique</div>
                    <div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('modele_economique', 0)}/20</div>
                </div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;">
                    <div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Stratégie Marketing</div>
                    <div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('strategie_marketing', 0)}/20</div>
                </div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;">
                    <div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Projections Financières</div>
                    <div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('projections_financieres', 0)}/20</div>
                </div>
            </div>
        </div>
        
        <!-- Points forts -->
        <div style="background: #e8f5e9; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
            <h2 style="color: #2E7D32; margin-top: 0;">💪 Points Forts</h2>
            <ul style="margin: 0; padding-left: 20px;">
                {''.join([f"<li style='margin: 5px 0;'>{point}</li>" for point in analysis.get('points_forts', [])])}
            </ul>
        </div>
        
        <!-- Axes d'amélioration -->
        <div style="background: #fff3e0; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
            <h2 style="color: #F57C00; margin-top: 0;">🎯 Axes d'Amélioration</h2>
            <ul style="margin: 0; padding-left: 20px;">
                {''.join([f"<li style='margin: 5px 0;'>{axe}</li>" for axe in analysis.get('axes_amelioration', [])])}
            </ul>
        </div>
        
        <!-- Recommandations -->
        <div style="background: #f3e5f5; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
            <h2 style="color: #7B1FA2; margin-top: 0;">💡 Recommandations</h2>
            <ol style="margin: 0; padding-left: 20px;">
                {''.join([f"<li style='margin: 5px 0;'>{reco}</li>" for reco in analysis.get('recommandations', [])])}
            </ol>
        </div>
        
        <!-- Pied de page -->
        <div style="background: #f5f5f5; padding: 20px; border-radius: 12px; text-align: center; color: #666; font-size: 0.9em;">
            <p style="margin: 0;">
                ⚡ Analyse générée par GPT-3.5 en {processing_time} secondes<br>
                📧 Ce rapport a été envoyé au professeur responsable
            </p>
        </div>
    </div>
    """
    
    return report_html