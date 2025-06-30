import httpx
import PyPDF2
import docx
import io
import os
import json
from datetime import datetime
from typing import Dict, Any
from openai import AsyncOpenAI

# ==========================================================
# 1. CONFIGURATION (Inchangée)
# ==========================================================
try:
    client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception as e:
    print(f"⚠️ Erreur initialisation OpenAI: {e}")
    import openai
    openai.api_key = os.getenv("OPENAI_API_KEY")

# ==========================================================
# 2. EXTRACTION DE TEXTE (Avec le correctif pour l'URL)
# ==========================================================
async def extract_text_from_file(file_url: str) -> str:
    """
    Télécharge un fichier depuis une URL (PDF ou DOCX), en extrait le texte,
    et applique une logique de troncature pour économiser les tokens.
    """
    print(f"INFO: Tentative de téléchargement et d'extraction depuis {file_url}")

    async with httpx.AsyncClient() as http_client:
        try:
            response = await http_client.get(file_url, follow_redirects=True, timeout=30.0)
            response.raise_for_status()
            file_content_bytes = response.content
        except httpx.RequestError as e:
            print(f"ERREUR: Échec du téléchargement du fichier: {e}")
            raise ValueError(f"Impossible de télécharger le fichier depuis l'URL: {file_url}")

    file_in_memory = io.BytesIO(file_content_bytes)
    text = ""

    # --- LE CORRECTIF EST ICI ---
    # On nettoie l'URL pour enlever les paramètres comme '?'
    clean_url = file_url.split('?')[0]

    try:
        if clean_url.lower().endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file_in_memory)
            num_pages = len(pdf_reader.pages)
            max_pages = min(num_pages, 10)
            
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            if num_pages > 10:
                text += f"\n[Note: Document tronqué - {num_pages} pages au total, les 10 premières pages ont été analysées]"
            print("INFO: Extraction PDF réussie.")

        elif clean_url.lower().endswith('.docx'):
            document = docx.Document(file_in_memory)
            max_paragraphs = 100
            paragraphs = document.paragraphs[:max_paragraphs]
            
            for paragraph in paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            if len(document.paragraphs) > max_paragraphs:
                text += f"\n[Note: Document tronqué - {len(document.paragraphs)} paragraphes au total, les 100 premiers ont été analysés]"
            print("INFO: Extraction DOCX réussie.")
        else:
            raise ValueError(f"Format de fichier non supporté dans l'URL nettoyée: {clean_url}")
            
    except Exception as e:
        print(f"ERREUR: Échec de l'analyse du contenu du fichier. {e}")
        raise ValueError(f"Erreur lors de l'analyse du contenu du fichier: {file_url}")

    return text

# ==========================================================
# 3. ANALYSE IA (MISE À JOUR AVEC VALIDATION)
# ==========================================================
async def analyze_business_plan(text: str, student_name: str, project_title: str) -> Dict[str, Any]:
    """Analyser un plan d'affaires avec GPT-3.5-turbo (économique)"""
    words = text.split()
    if len(words) > 3000:
        text = ' '.join(words[:3000]) + "\n\n[Document tronqué pour l'analyse]"
    
    # NOUVEAU PROMPT AMÉLIORÉ
    system_prompt = """Tu es un expert en évaluation de plans d'affaires académiques.

ÉTAPE 1 - VALIDATION : Vérifie d'abord si le document est un plan d'affaires valide.
Un plan d'affaires DOIT contenir au minimum :
- Une description du produit/service
- Une analyse de marché ou clientèle cible
- Un modèle économique ou stratégie de revenus

Si le document n'est PAS un plan d'affaires (ex: recette, histoire, devoir non pertinent), retourne :
{
    "document_valide": false,
    "raison_rejet": "Ce document n'est pas un plan d'affaires. [Explique pourquoi]",
    "score_global": 0
}

ÉTAPE 2 - ANALYSE : Si c'est un plan d'affaires valide, analyse-le et retourne :
{
    "document_valide": true,
    "resume_executif": "résumé objectif en 2-3 phrases",
    "score_global": nombre entre 0 et 100 (sois exigeant, moyenne = 50-60),
    "scores": {
        "viabilite_concept": nombre entre 0 et 20 (0 si absent, note selon qualité),
        "etude_marche": nombre entre 0 et 20 (0 si aucune analyse marché),
        "modele_economique": nombre entre 0 et 20 (0 si aucun modèle de revenus),
        "strategie_marketing": nombre entre 0 et 20 (0 si absente),
        "projections_financieres": nombre entre 0 et 20 (0 si absentes)
    },
    "completude": "pourcentage des sections présentes (0-100%)",
    "points_forts": ["maximum 3 points spécifiques au document"],
    "axes_amelioration": ["maximum 3 axes concrets"],
    "recommandations": ["maximum 3 actions prioritaires"]
}

IMPORTANT : 
- Sois STRICT sur les scores. Un document incomplet < 40/100
- Note 0 les sections totalement absentes
- Retourne UNIQUEMENT le JSON, aucun autre texte"""
    
    user_prompt = f"Plan d'affaires de {student_name} - Projet: {project_title}\n\n{text}\n\nFournis l'analyse JSON."
    
    try:
        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000,
            response_format={"type": "json_object"}
        )
        analysis = json.loads(response.choices[0].message.content)
        
        # Vérifier si le document est valide
        if not analysis.get('document_valide', True):
            # Document rejeté
            return {
                "document_valide": False,
                "raison_rejet": analysis.get('raison_rejet', 'Document non conforme'),
                "score_global": 0,
                "scores": {
                    "viabilite_concept": 0,
                    "etude_marche": 0,
                    "modele_economique": 0,
                    "strategie_marketing": 0,
                    "projections_financieres": 0
                },
                "resume_executif": analysis.get('raison_rejet', 'Document rejeté'),
                "points_forts": [],
                "axes_amelioration": ["Document non conforme aux exigences d'un plan d'affaires"],
                "recommandations": ["Soumettre un véritable plan d'affaires"]
            }
        
        # Document valide - continuer avec le traitement normal
        if 'scores' not in analysis: 
            analysis['scores'] = {}
        
        default_scores = {
            'viabilite_concept': 10, 
            'etude_marche': 10, 
            'modele_economique': 10, 
            'strategie_marketing': 10, 
            'projections_financieres': 10
        }
        
        for key, default in default_scores.items():
            if key not in analysis['scores']: 
                analysis['scores'][key] = default
        
        # Convertir les scores en entiers
        for key, value in analysis['scores'].items():
            try:
                analysis['scores'][key] = int(value)
            except (ValueError, TypeError):
                analysis['scores'][key] = 0  # 0 si erreur de conversion

        # Calculer le score global si absent
        if 'score_global' not in analysis:
            total = sum(analysis['scores'].values())
            analysis['score_global'] = int(total)
        else:
            try:
                analysis['score_global'] = int(analysis['score_global'])
            except:
                analysis['score_global'] = sum(analysis['scores'].values())

        return analysis
        
    except Exception as e:
        print(f"Erreur OpenAI: {e}")
        return {
            "document_valide": True,
            "error": str(e), 
            "resume_executif": "Erreur lors de l'analyse automatique.", 
            "score_global": 50, 
            "scores": default_scores, 
            "points_forts": ["Analyse non disponible"], 
            "axes_amelioration": ["Analyse non disponible"], 
            "recommandations": ["Réessayer l'analyse"]
        }

# ==========================================================
# 4. GÉNÉRATION DU RAPPORT (MISE À JOUR POUR GÉRER LES REJETS)
# ==========================================================
def generate_formatted_report(analysis: Dict[str, Any], student_name: str, project_title: str, processing_time: int) -> str:
    """Générer un rapport HTML formaté"""
    
    # Vérifier si le document a été rejeté
    if not analysis.get('document_valide', True):
        return f"""
        <div style="font-family: 'Inter', -apple-system, sans-serif; line-height: 1.6; color: #333;">
            <div style="background: linear-gradient(135deg, #f44336 0%, #d32f2f 100%); color: white; padding: 30px; border-radius: 16px; margin-bottom: 30px; text-align: center;">
                <h1 style="margin: 0 0 10px 0;">❌ Document Non Valide</h1>
                <p style="margin: 0; opacity: 0.9;">Ce document n'est pas un plan d'affaires</p>
            </div>
            <div style="background: #ffebee; padding: 25px; border-radius: 12px; margin-bottom: 25px; border-left: 4px solid #f44336;">
                <h2 style="color: #c62828; margin-top: 0;">🚫 Raison du Rejet</h2>
                <p style="margin: 0; font-size: 1.1em;">{analysis.get('raison_rejet', 'Document non conforme aux exigences')}</p>
            </div>
            <div style="background: #f8f9fa; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
                <h2 style="color: #333; margin-top: 0;">📋 Informations de la Soumission</h2>
                <table style="width: 100%; border-collapse: collapse;">
                    <tr><td style="padding: 8px 0;"><strong>Étudiant:</strong></td><td style="padding: 8px 0;">{student_name}</td></tr>
                    <tr><td style="padding: 8px 0;"><strong>Titre du projet:</strong></td><td style="padding: 8px 0;">{project_title}</td></tr>
                    <tr><td style="padding: 8px 0;"><strong>Date:</strong></td><td style="padding: 8px 0;">{datetime.now().strftime('%d/%m/%Y à %H:%M')}</td></tr>
                    <tr><td style="padding: 8px 0;"><strong>Score:</strong></td><td style="padding: 8px 0;"><span style="font-size: 1.5em; color: #f44336; font-weight: bold;">0/100</span></td></tr>
                </table>
            </div>
            <div style="background: #fff3e0; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
                <h2 style="color: #F57C00; margin-top: 0;">⚠️ Recommandation</h2>
                <p style="margin: 0;">L'étudiant doit soumettre un véritable plan d'affaires comprenant au minimum :</p>
                <ul style="margin: 10px 0 0 20px;">
                    <li>Une description claire du produit ou service</li>
                    <li>Une analyse du marché cible</li>
                    <li>Un modèle économique ou stratégie de revenus</li>
                </ul>
            </div>
            <div style="background: #f5f5f5; padding: 20px; border-radius: 12px; text-align: center; color: #666; font-size: 0.9em;">
                <p style="margin: 0;">⚡ Analyse effectuée en {processing_time} secondes<br>📧 Notification envoyée au professeur</p>
            </div>
        </div>
        """
    
    # Rapport normal pour un document valide
    scores = analysis.get('scores', {})
    score_global = analysis.get('score_global', 0)
    
    if score_global >= 80: color_score = "#4CAF50"
    elif score_global >= 60: color_score = "#FF9800"
    else: color_score = "#f44336"
    
    completude = analysis.get('completude', 'N/A')
    
    report_html = f"""
    <div style="font-family: 'Inter', -apple-system, sans-serif; line-height: 1.6; color: #333;">
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 16px; margin-bottom: 30px;">
            <h1 style="margin: 0 0 10px 0;">📊 Rapport d'Analyse Automatique</h1>
            <p style="margin: 0; opacity: 0.9;">Évaluation par Intelligence Artificielle</p>
        </div>
        <div style="background: #f8f9fa; padding: 25px; border-radius: 12px; margin-bottom: 25px;">
            <h2 style="color: #333; margin-top: 0;">📋 Informations du Projet</h2>
            <table style="width: 100%; border-collapse: collapse;">
                <tr><td style="padding: 8px 0;"><strong>Étudiant:</strong></td><td style="padding: 8px 0;">{student_name}</td></tr>
                <tr><td style="padding: 8px 0;"><strong>Projet:</strong></td><td style="padding: 8px 0;">{project_title}</td></tr>
                <tr><td style="padding: 8px 0;"><strong>Date:</strong></td><td style="padding: 8px 0;">{datetime.now().strftime('%d/%m/%Y à %H:%M')}</td></tr>
                <tr><td style="padding: 8px 0;"><strong>Score Global:</strong></td><td style="padding: 8px 0;"><span style="font-size: 1.5em; color: {color_score}; font-weight: bold;">{score_global}/100</span></td></tr>
                <tr><td style="padding: 8px 0;"><strong>Complétude:</strong></td><td style="padding: 8px 0;">{completude}</td></tr>
            </table>
        </div>
        <div style="background: #e3f2fd; padding: 25px; border-radius: 12px; margin-bottom: 25px; border-left: 4px solid #2196F3;">
            <h2 style="color: #1976D2; margin-top: 0;">🎯 Résumé Exécutif</h2>
            <p style="margin: 0;">{analysis.get('resume_executif', 'Non disponible')}</p>
        </div>
        <div style="margin-bottom: 25px;">
            <h2 style="color: #333;">📊 Scores Détaillés</h2>
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 15px;">
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;"><div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Viabilité du Concept</div><div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('viabilite_concept', 0)}/20</div></div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;"><div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Étude de Marché</div><div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('etude_marche', 0)}/20</div></div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;"><div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Modèle Économique</div><div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('modele_economique', 0)}/20</div></div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;"><div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Stratégie Marketing</div><div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('strategie_marketing', 0)}/20</div></div>
                <div style="background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;"><div style="color: #666; font-size: 0.9em; margin-bottom: 5px;">Projections Financières</div><div style="font-size: 1.8em; font-weight: bold; color: #667eea;">{scores.get('projections_financieres', 0)}/20</div></div>
            </div>
        </div>
        <div style="background: #e8f5e9; padding: 25px; border-radius: 12px; margin-bottom: 25px;"><h2 style="color: #2E7D32; margin-top: 0;">💪 Points Forts</h2><ul style="margin: 0; padding-left: 20px;">{''.join([f"<li style='margin: 5px 0;'>{point}</li>" for point in analysis.get('points_forts', [])])}</ul></div>
        <div style="background: #fff3e0; padding: 25px; border-radius: 12px; margin-bottom: 25px;"><h2 style="color: #F57C00; margin-top: 0;">🎯 Axes d'Amélioration</h2><ul style="margin: 0; padding-left: 20px;">{''.join([f"<li style='margin: 5px 0;'>{axe}</li>" for axe in analysis.get('axes_amelioration', [])])}</ul></div>
        <div style="background: #f3e5f5; padding: 25px; border-radius: 12px; margin-bottom: 25px;"><h2 style="color: #7B1FA2; margin-top: 0;">💡 Recommandations</h2><ol style="margin: 0; padding-left: 20px;">{''.join([f"<li style='margin: 5px 0;'>{reco}</li>" for reco in analysis.get('recommandations', [])])}</ol></div>
        <div style="background: #f5f5f5; padding: 20px; border-radius: 12px; text-align: center; color: #666; font-size: 0.9em;"><p style="margin: 0;">⚡ Analyse générée par GPT-3.5 en {processing_time} secondes<br>📧 Ce rapport a été envoyé au professeur responsable</p></div>
    </div>
    """
    return report_html